"""Export cover letter to .docx using python-docx."""
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def cover_letter_to_docx(cover_letter_text: str, applicant_name: str = "", job_title: str = "", company: str = "") -> bytes:
    """
    Convert cover letter text to a formatted .docx file.
    Returns bytes of the .docx file.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Header with applicant name
    if applicant_name:
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run(applicant_name)
        run.bold = True
        run.font.size = Pt(14)

    # Subject line
    if job_title and company:
        subject = doc.add_paragraph()
        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subj_run = subject.add_run(f"Application for {job_title} at {company}")
        subj_run.font.size = Pt(11)
        subj_run.italic = True

    doc.add_paragraph()  # Spacer

    # Cover letter body — split by paragraphs
    paragraphs = cover_letter_text.strip().split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = para.add_run(para_text.replace("\n", " "))
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(8)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def resume_to_docx(resume_text: str, applicant_name: str = "") -> bytes:
    """Convert rewritten resume text to a .docx file."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    if applicant_name:
        h = doc.add_heading(applicant_name, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in resume_text.strip().split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Detect section headers (ALL CAPS or short lines ending with :)
        if line.isupper() or (len(line) < 40 and line.endswith(":")):
            para = doc.add_heading(line.rstrip(":"), level=2)
        elif line.startswith(("•", "-", "*", "▪", "◦")):
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(line.lstrip("•-*▪◦ "))
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
